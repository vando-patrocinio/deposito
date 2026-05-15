"""Budget AI — Orçamento_IA.

Endpoints REST para o módulo Orçamento. Fluxo:
  1) POST /api/budget                       — cria orçamento (nome + descrição)
  2) POST /api/budget/{id}/upload-csv       — upload CSV de itens (item;qtde;unidade;spec)
  3) POST /api/budget/{id}/analyze          — Orçamento_IA estima 3 preços/item + média
  4) PUT  /api/budget/{id}                  — edita %ganho · %imposto · %mão-de-obra · overrides
  5) GET  /api/budget/{id}                  — detalhes (com cálculos)
  6) GET  /api/budget                       — lista
  7) GET  /api/budget/{id}/pdf              — romaneio imprimível
  8) GET  /api/budget/kpis                  — KPIs do módulo

Acesso: administrador · gestor · financeiro.
"""
from __future__ import annotations

import csv
import io
import json
import logging
import re
import uuid
from datetime import datetime, timezone
from typing import Any, Dict, List, Optional

from fastapi import APIRouter, Depends, File, HTTPException, Query, UploadFile
from fastapi.responses import Response
from pydantic import BaseModel, Field

from core import DEMO_COMPANY_ID, now_iso, require_role
from database import db
from services.motor_ia import chat_completion

logger = logging.getLogger("ponto.budget")
router = APIRouter(prefix="/api/budget", tags=["budget"])

# ---------------------------------------------------------------------------
# Models
# ---------------------------------------------------------------------------


class BudgetCreateIn(BaseModel):
    name: str = Field(..., min_length=2, max_length=200)
    description: Optional[str] = None


class BudgetUpdateIn(BaseModel):
    name: Optional[str] = None
    description: Optional[str] = None
    margin_pct: Optional[float] = Field(None, ge=0, le=500)
    tax_pct: Optional[float] = Field(None, ge=0, le=100)
    labor_pct: Optional[float] = Field(None, ge=0, le=500)
    items: Optional[List[Dict[str, Any]]] = None  # overrides de avg/manual


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _new_id(prefix: str) -> str:
    return f"{prefix}-{uuid.uuid4().hex[:10]}"


def _money(v: float) -> float:
    return round(float(v or 0), 2)


def _resolve_unit_price(item: Dict[str, Any]) -> float:
    """Resolve o preço unitário aplicado, na ordem de prioridade:
    1. manual_override (se setado e != "")
    2. preço escolhido via price_choice (low/mid/high) buscado em `prices`
    3. avg_price (fallback — compat com itens antigos sem price_choice)
    """
    if item.get("manual_override") not in (None, ""):
        try:
            return float(item["manual_override"])
        except (ValueError, TypeError):
            pass
    choice = (item.get("price_choice") or "").lower()
    prices = item.get("prices") or []
    # prices vem como [{"label":"Baixo","value":X},{"label":"Médio",...},...]
    label_map = {"low": "baixo", "mid": "médio", "high": "alto",
                  "baixo": "baixo", "médio": "médio", "medio": "médio",
                  "alto": "alto"}
    target = label_map.get(choice)
    if target:
        for p in prices:
            if (p.get("label") or "").strip().lower() == target:
                try:
                    return float(p.get("value") or 0)
                except (ValueError, TypeError):
                    pass
    try:
        return float(item.get("avg_price") or 0)
    except (ValueError, TypeError):
        return 0.0


def _calc_totals(items: List[Dict[str, Any]], margin_pct: float,
                   tax_pct: float, labor_pct: float) -> Dict[str, float]:
    """Calcula totais. base = soma(qtde * unit). unit obedece a price_choice
    do item (low/mid/high) ou manual_override quando setado. Final = base *
    (1+ganho+labor+imposto).
    """
    base = 0.0
    for it in items or []:
        qty = float(it.get("qty") or 0)
        unit = _resolve_unit_price(it)
        base += qty * unit
    base = _money(base)
    margin_val = _money(base * (margin_pct or 0) / 100.0)
    labor_val = _money(base * (labor_pct or 0) / 100.0)
    subtotal = _money(base + margin_val + labor_val)
    tax_val = _money(subtotal * (tax_pct or 0) / 100.0)
    final = _money(subtotal + tax_val)
    return {
        "base": base, "margin_val": margin_val, "labor_val": labor_val,
        "subtotal": subtotal, "tax_val": tax_val, "final": final,
    }


async def _get_budget(cid: str, bid: str) -> Dict[str, Any]:
    doc = await db.budgets.find_one({"id": bid, "company_id": cid}, {"_id": 0})
    if not doc:
        raise HTTPException(404, "Orçamento não encontrado")
    return doc


def _hydrate(doc: Dict[str, Any]) -> Dict[str, Any]:
    """Adiciona totais calculados ao doc antes de devolver."""
    items = doc.get("items") or []
    totals = _calc_totals(items, doc.get("margin_pct", 25),
                            doc.get("tax_pct", 0), doc.get("labor_pct", 0))
    doc["totals"] = totals
    return doc


# ---------------------------------------------------------------------------
# CRUD
# ---------------------------------------------------------------------------


@router.post("")
async def create_budget(body: BudgetCreateIn,
                          user: dict = Depends(require_role("administrador",
                                                              "gestor",
                                                              "financeiro"))):
    cid = user.get("company_id") or DEMO_COMPANY_ID
    doc = {
        "id": _new_id("bud"),
        "company_id": cid,
        "name": body.name.strip(),
        "description": (body.description or "").strip(),
        "status": "draft",  # draft → analyzed → final
        "items": [],
        "margin_pct": 25.0,
        "tax_pct": 0.0,
        "labor_pct": 0.0,
        "created_by_user_id": user.get("id"),
        "created_by_name": user.get("name"),
        "created_at": now_iso(),
        "updated_at": now_iso(),
    }
    await db.budgets.insert_one(doc.copy())
    doc.pop("_id", None)
    return _hydrate(doc)


@router.get("")
async def list_budgets(user: dict = Depends(require_role("administrador",
                                                            "gestor",
                                                            "financeiro"))):
    cid = user.get("company_id") or DEMO_COMPANY_ID
    items = await db.budgets.find(
        {"company_id": cid}, {"_id": 0},
    ).sort("created_at", -1).to_list(200)
    return {"items": [_hydrate(i) for i in items], "total": len(items)}


@router.get("/kpis")
async def budget_kpis(user: dict = Depends(require_role("administrador",
                                                            "gestor",
                                                            "financeiro"))):
    cid = user.get("company_id") or DEMO_COMPANY_ID
    docs = await db.budgets.find(
        {"company_id": cid}, {"_id": 0, "items": 1, "margin_pct": 1,
                                "tax_pct": 1, "labor_pct": 1, "status": 1,
                                "created_at": 1},
    ).to_list(500)
    total = len(docs)
    draft = sum(1 for d in docs if d.get("status") == "draft")
    analyzed = sum(1 for d in docs if d.get("status") == "analyzed")
    final_status = sum(1 for d in docs if d.get("status") == "final")
    margins = [d.get("margin_pct", 0) for d in docs if d.get("margin_pct")]
    total_final_value = 0.0
    for d in docs:
        t = _calc_totals(d.get("items") or [], d.get("margin_pct", 0),
                          d.get("tax_pct", 0), d.get("labor_pct", 0))
        total_final_value += t["final"]
    return {
        "total": total,
        "draft": draft,
        "analyzed": analyzed,
        "final": final_status,
        "avg_margin_pct": round(sum(margins) / len(margins), 1) if margins else 0,
        "total_value": _money(total_final_value),
        "avg_value": _money(total_final_value / total) if total else 0,
    }


@router.get("/{bid}")
async def get_budget(bid: str,
                       user: dict = Depends(require_role("administrador",
                                                           "gestor",
                                                           "financeiro"))):
    cid = user.get("company_id") or DEMO_COMPANY_ID
    return _hydrate(await _get_budget(cid, bid))


@router.put("/{bid}")
async def update_budget(bid: str, body: BudgetUpdateIn,
                          user: dict = Depends(require_role("administrador",
                                                              "gestor",
                                                              "financeiro"))):
    cid = user.get("company_id") or DEMO_COMPANY_ID
    doc = await _get_budget(cid, bid)
    upd: Dict[str, Any] = {"updated_at": now_iso()}
    if body.name is not None:
        upd["name"] = body.name.strip()
    if body.description is not None:
        upd["description"] = body.description.strip()
    if body.margin_pct is not None:
        upd["margin_pct"] = float(body.margin_pct)
    if body.tax_pct is not None:
        upd["tax_pct"] = float(body.tax_pct)
    if body.labor_pct is not None:
        upd["labor_pct"] = float(body.labor_pct)
    if body.items is not None:
        # Mantém os campos originais e aplica overrides nos campos editáveis.
        # manual_override aceita None (= limpar override e voltar ao price_choice).
        existing = {it["id"]: it for it in doc.get("items", []) if it.get("id")}
        new_items: List[Dict[str, Any]] = []
        for it in body.items:
            iid = it.get("id")
            base = existing.get(iid, {})
            # Aceita explicitamente manual_override=None (limpar) mas filtra
            # outros valores None para não zerar dados acidentalmente.
            patch = {k: v for k, v in it.items()
                       if v is not None or k == "manual_override"}
            merged = {**base, **patch}
            new_items.append(merged)
        upd["items"] = new_items
    await db.budgets.update_one({"id": bid, "company_id": cid}, {"$set": upd})
    return _hydrate(await _get_budget(cid, bid))


@router.delete("/{bid}")
async def delete_budget(bid: str,
                          user: dict = Depends(require_role("administrador",
                                                              "gestor"))):
    cid = user.get("company_id") or DEMO_COMPANY_ID
    res = await db.budgets.delete_one({"id": bid, "company_id": cid})
    if res.deleted_count == 0:
        raise HTTPException(404, "Orçamento não encontrado")
    return {"ok": True}


# ---------------------------------------------------------------------------
# Upload CSV
# ---------------------------------------------------------------------------


@router.post("/{bid}/upload-csv")
async def upload_csv(bid: str, file: UploadFile = File(...),
                       user: dict = Depends(require_role("administrador",
                                                           "gestor",
                                                           "financeiro"))):
    """Compat: endpoint legado redireciona para o parser genérico.

    Aceita CSV / PDF / DOCX — detecção automática pelo content-type/extensão.
    """
    return await upload_file(bid, file, user)


@router.post("/{bid}/upload")
async def upload_file(bid: str, file: UploadFile = File(...),
                        user: dict = Depends(require_role("administrador",
                                                            "gestor",
                                                            "financeiro"))):
    """Sobe arquivo de itens — suporta CSV, PDF ou DOCX.

    Fluxo:
    - CSV: parser direto (colunas item·qtde·unidade·especificacao).
    - PDF/DOCX: extrai texto bruto + manda pra Orçamento_IA estruturar em
      itens JSON (mesma chamada Claude usada no `/analyze`).

    Retorna `{ok, items_count, source}` onde `source` indica de qual parser
    veio (csv·pdf·docx).
    """
    cid = user.get("company_id") or DEMO_COMPANY_ID
    await _get_budget(cid, bid)
    raw = await file.read()
    if not raw:
        raise HTTPException(400, "Arquivo vazio")
    if len(raw) > 5 * 1024 * 1024:
        raise HTTPException(413, "Arquivo > 5MB")

    fname = (file.filename or "").lower()
    ctype = (file.content_type or "").lower()

    items: List[Dict[str, Any]] = []
    source = "csv"

    if fname.endswith(".csv") or "csv" in ctype or "text/plain" in ctype:
        items = _parse_csv_bytes(raw)
        source = "csv"
    elif fname.endswith(".pdf") or "pdf" in ctype:
        text = _extract_pdf_text(raw)
        items = await _extract_items_via_ai(cid, text, source="pdf")
        source = "pdf"
    elif fname.endswith(".docx") or "officedocument.wordprocessingml" in ctype:
        text = _extract_docx_text(raw)
        items = await _extract_items_via_ai(cid, text, source="docx")
        source = "docx"
    elif fname.endswith(".doc"):
        raise HTTPException(415, "Formato .doc antigo não suportado — salve como "
                                  ".docx ou .pdf no Word/LibreOffice e tente novamente.")
    else:
        raise HTTPException(415, f"Formato não suportado ({ctype or fname}). "
                                   "Use CSV, PDF ou DOCX.")

    if not items:
        raise HTTPException(400, "Nenhum item válido extraído do arquivo.")

    await db.budgets.update_one(
        {"id": bid, "company_id": cid},
        {"$set": {"items": items, "status": "draft", "source_file": fname,
                   "source_type": source, "updated_at": now_iso()}},
    )
    return {"ok": True, "items_count": len(items), "source": source}


# ---------------------------------------------------------------------------
# File parsers
# ---------------------------------------------------------------------------


def _parse_csv_bytes(raw: bytes) -> List[Dict[str, Any]]:
    """Parser de CSV com colunas item;qtde;unidade;especificacao."""
    try:
        text = raw.decode("utf-8-sig")
    except UnicodeDecodeError:
        try:
            text = raw.decode("latin-1")
        except UnicodeDecodeError:
            raise HTTPException(400, "Encoding não suportado (use UTF-8).")
    sample = text[:1024]
    sep = ";" if sample.count(";") >= sample.count(",") else ","
    reader = csv.DictReader(io.StringIO(text), delimiter=sep)
    items: List[Dict[str, Any]] = []
    for row in reader:
        if not row:
            continue
        keys = {(k or "").strip().lower(): v for k, v in row.items()}
        name = (keys.get("item") or keys.get("produto") or keys.get("material")
                 or keys.get("descricao") or keys.get("descrição") or "").strip()
        if not name:
            continue
        qty_raw = (keys.get("qtde") or keys.get("qty") or keys.get("quantidade")
                    or keys.get("quant") or "1").strip()
        unit = (keys.get("unidade") or keys.get("un") or keys.get("unit")
                 or "un").strip() or "un"
        spec = (keys.get("especificacao") or keys.get("especificação")
                 or keys.get("spec") or keys.get("observacao") or "").strip()
        try:
            qty = float(re.sub(r"[^\d.,-]", "", qty_raw).replace(",", "."))
        except (ValueError, TypeError):
            qty = 1.0
        items.append({
            "id": _new_id("itm"),
            "name": name, "qty": qty, "unit": unit, "spec": spec,
            "prices": [], "avg_price": 0.0, "manual_override": None,
        })
    return items


def _extract_pdf_text(raw: bytes) -> str:
    """Extrai texto de PDF usando pdfplumber."""
    import pdfplumber
    out_lines: List[str] = []
    try:
        with pdfplumber.open(io.BytesIO(raw)) as pdf:
            for page in pdf.pages:
                t = page.extract_text() or ""
                if t:
                    out_lines.append(t)
                # Tenta também tabelas (mais preciso pra orçamentos formatados)
                for tbl in (page.extract_tables() or []):
                    for row in tbl or []:
                        line = " | ".join(
                            (c or "").strip() for c in row if c)
                        if line.strip():
                            out_lines.append(line)
    except Exception as e:
        logger.warning("[budget] PDF extract falhou: %s", e)
        raise HTTPException(422, f"Falha ao ler PDF: {e}")
    text = "\n".join(out_lines).strip()
    if not text:
        raise HTTPException(422, "PDF sem texto extraível (talvez seja imagem "
                                  "escaneada). Tente OCR ou converta para "
                                  "DOCX/CSV.")
    return text[:30000]  # cap para não estourar contexto da IA


def _extract_docx_text(raw: bytes) -> str:
    """Extrai texto de DOCX usando python-docx (paragraphs + tables)."""
    from docx import Document
    out_lines: List[str] = []
    try:
        doc_obj = Document(io.BytesIO(raw))
        for p in doc_obj.paragraphs:
            t = (p.text or "").strip()
            if t:
                out_lines.append(t)
        for tbl in doc_obj.tables:
            for row in tbl.rows:
                cells = [c.text.strip() for c in row.cells]
                line = " | ".join(c for c in cells if c)
                if line.strip():
                    out_lines.append(line)
    except Exception as e:
        logger.warning("[budget] DOCX extract falhou: %s", e)
        raise HTTPException(422, f"Falha ao ler DOCX: {e}")
    text = "\n".join(out_lines).strip()
    if not text:
        raise HTTPException(422, "DOCX vazio ou sem texto extraível.")
    return text[:30000]


async def _extract_items_via_ai(cid: str, text: str,
                                  source: str = "pdf") -> List[Dict[str, Any]]:
    """Usa Claude para extrair lista estruturada de itens a partir de texto
    bruto (PDF ou DOCX). Retorna lista no MESMO formato do CSV parser."""
    if not text:
        return []
    sys_prompt = (
        "Você é Orçamento_IA, extrator de itens de orçamentos para "
        "provedores de internet. Recebe texto extraído de um documento "
        f"({source.upper()}) e devolve a lista de itens em JSON puro. "
        "Identifique nome do item, quantidade, unidade (un/m/rolo/cx/kg) "
        "e especificação técnica. Ignore cabeçalhos, totais, assinaturas e "
        "comentários. Se quantidade não for explícita, use 1. Se unidade "
        "não for explícita, use 'un'."
    )
    user_prompt = (
        "Extraia os itens em JSON estrito no formato:\n"
        '{"items":[{"name":"...","qty":2,"unit":"un","spec":"..."}]}\n\n'
        f"TEXTO DO DOCUMENTO ({source.upper()}):\n{text}"
    )
    try:
        result = await chat_completion(
            cid,
            messages=[
                {"role": "system", "content": sys_prompt},
                {"role": "user", "content": user_prompt},
            ],
            temperature=0.2,
            max_tokens=6000,
            purpose="general",
            agent="orcamento_ai",
        )
        content = (result.get("content") or "").strip()
    except Exception as e:
        logger.exception("[budget] AI extract falhou: %s", e)
        raise HTTPException(503, f"Orçamento_IA indisponível: {e}")

    m = re.search(r"\{[\s\S]*\}", content)
    if not m:
        raise HTTPException(422, f"IA não retornou itens estruturados: "
                                   f"{content[:200]}")
    try:
        parsed = json.loads(m.group(0))
    except json.JSONDecodeError as e:
        raise HTTPException(422, f"JSON inválido da IA: {e}")

    items: List[Dict[str, Any]] = []
    for it in parsed.get("items", []):
        name = (it.get("name") or "").strip()
        if not name:
            continue
        try:
            qty = float(it.get("qty") or 1)
        except (ValueError, TypeError):
            qty = 1.0
        items.append({
            "id": _new_id("itm"),
            "name": name,
            "qty": qty,
            "unit": (it.get("unit") or "un").strip(),
            "spec": (it.get("spec") or "").strip(),
            "prices": [], "avg_price": 0.0, "manual_override": None,
        })
    return items


# ---------------------------------------------------------------------------
# IA — estima 3 preços + média
# ---------------------------------------------------------------------------


@router.post("/{bid}/analyze")
async def analyze_budget(bid: str,
                           user: dict = Depends(require_role("administrador",
                                                               "gestor",
                                                               "financeiro"))):
    """Orçamento_IA estima 3 preços por item baseado em conhecimento do
    mercado brasileiro (Claude Sonnet 4.5 via Emergent LLM Key). Calcula
    média e popula `avg_price` em cada item.

    NOTA: estimativas pontuais (sem web scraping). Para preços live, o
    usuário pode editar manualmente o campo `manual_override` na revisão.
    """
    cid = user.get("company_id") or DEMO_COMPANY_ID
    doc = await _get_budget(cid, bid)
    items = doc.get("items") or []
    if not items:
        raise HTTPException(400, "Orçamento sem itens — suba o CSV primeiro")

    # Monta prompt para Claude
    item_lines = "\n".join([
        f"- ID:{it['id']} · {it['name']} · qtde={it['qty']} {it['unit']}"
        f"{(' · spec=' + it['spec']) if it.get('spec') else ''}"
        for it in items
    ])
    sys_prompt = (
        "Você é Orçamento_IA, especialista em precificação de materiais para "
        "provedores de internet (ISPs) no Brasil. Recebe uma lista de itens "
        "e retorna 3 estimativas de preço por item (baixo/médio/alto) baseado "
        "no seu conhecimento dos sites Mercado Livre, Amazon, Magazine Luiza, "
        "Furukawa, FiberHome, Intelbras, Mikrotik, Datacom. Sempre em REAIS "
        "(R$). Considere o preço para 1 unidade do item (não multiplique pela "
        "quantidade — o sistema faz isso). Se item desconhecido, use estimativa "
        "razoável e marque confidence='low'."
    )
    user_prompt = (
        f"Estime 3 preços (baixo/médio/alto) por item. Retorne APENAS JSON "
        f"válido no formato:\n"
        f'{{"items":[{{"id":"itm-xxxx","low":12.50,"mid":15.00,"high":18.90,'
        f'"sources":["Mercado Livre","Furukawa","Amazon"],"confidence":"high"}}]}}\n\n'
        f"ITENS:\n{item_lines}"
    )

    try:
        result = await chat_completion(
            cid,
            messages=[
                {"role": "system", "content": sys_prompt},
                {"role": "user", "content": user_prompt},
            ],
            temperature=0.3,
            max_tokens=8000,
            purpose="general",
            agent="orcamento_ai",
        )
        content = (result.get("content") or "").strip()
    except Exception as e:
        logger.exception("[budget] Claude falhou: %s", e)
        raise HTTPException(503, f"Orçamento_IA indisponível: {e}")

    # Tenta extrair JSON (Claude às vezes vem com ```json ... ```)
    m = re.search(r"\{[\s\S]*\}", content)
    if not m:
        raise HTTPException(502, f"IA retornou formato inválido: {content[:200]}")
    try:
        parsed = json.loads(m.group(0))
    except json.JSONDecodeError as e:
        raise HTTPException(502, f"JSON inválido da IA: {e}")

    ai_items = {it.get("id"): it for it in parsed.get("items", []) if it.get("id")}
    # Atualiza cada item com prices + avg
    updated = []
    for it in items:
        ai = ai_items.get(it["id"])
        if ai:
            low = float(ai.get("low") or 0)
            mid = float(ai.get("mid") or 0)
            high = float(ai.get("high") or 0)
            avg = round((low + mid + high) / 3.0, 2) if (low + mid + high) else 0
            it["prices"] = [
                {"label": "Baixo", "value": _money(low)},
                {"label": "Médio", "value": _money(mid)},
                {"label": "Alto", "value": _money(high)},
            ]
            it["sources"] = ai.get("sources") or []
            it["confidence"] = ai.get("confidence") or "medium"
            it["avg_price"] = _money(avg)
            # Default escolha = "mid" (médio). Usuário pode trocar pra low/high
            # clicando nos chips na UI.
            if not it.get("price_choice"):
                it["price_choice"] = "mid"
        updated.append(it)

    await db.budgets.update_one(
        {"id": bid, "company_id": cid},
        {"$set": {
            "items": updated,
            "status": "analyzed",
            "analyzed_at": now_iso(),
            "analyzed_by": user.get("name"),
            "ai_model": result.get("model", ""),
            "updated_at": now_iso(),
        }},
    )
    return _hydrate(await _get_budget(cid, bid))


# ---------------------------------------------------------------------------
# PDF — romaneio imprimível
# ---------------------------------------------------------------------------


@router.get("/{bid}/pdf")
async def budget_pdf(bid: str,
                       user: dict = Depends(require_role("administrador",
                                                           "gestor",
                                                           "financeiro"))):
    cid = user.get("company_id") or DEMO_COMPANY_ID
    doc = _hydrate(await _get_budget(cid, bid))
    # Carrega branding (logo + dados da empresa) — mesmo cabeçalho do romaneio
    from routes.branding import get_branding
    from services.budget_pdf import build_budget_pdf
    branding_obj = await get_branding(cid)
    branding_dict = branding_obj.model_dump() if hasattr(branding_obj, "model_dump") \
                    else dict(branding_obj or {})
    pdf_bytes = build_budget_pdf(doc, generated_by=user.get("name"),
                                    branding=branding_dict)
    # Marca como final na primeira vez que gera o PDF
    if doc.get("status") != "final":
        await db.budgets.update_one(
            {"id": bid, "company_id": cid},
            {"$set": {"status": "final", "finalized_at": now_iso(),
                       "finalized_by": user.get("name")}},
        )
    safe = re.sub(r"[^\w\-]+", "-", doc.get("name", "orcamento")).strip("-")
    return Response(content=pdf_bytes, media_type="application/pdf", headers={
        "Content-Disposition": f"inline; filename=\"Orcamento-{safe}.pdf\"",
    })
